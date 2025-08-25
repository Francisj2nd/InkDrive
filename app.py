import os
import io
import re
import uuid
import requests
import markdown
import secrets
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, flash, session, abort
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from google.oauth2 import id_token
from google.auth.transport import requests as google_requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from google.api_core import exceptions as google_exceptions
import vertexai
from vertexai.generative_models import GenerativeModel
from datetime import datetime, timedelta
import json
import logging
from sqlalchemy.exc import OperationalError, DatabaseError
from sqlalchemy import func
import time
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# Import our models and forms
from models import db, User, Article, ChatSession
from forms import LoginForm, RegisterForm, ProfileForm, ChangePasswordForm

# Import admin blueprint
from admin import admin_bp

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- 1. INITIALIZATION & HELPERS ---
app = Flask(__name__)

# Register admin blueprint
app.register_blueprint(admin_bp)

# Configuration
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', secrets.token_hex(32))
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///inkdrive.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
# Engine options - conditionally add connect_args for PostgreSQL
engine_options = {
    'pool_pre_ping': True,
    'pool_recycle': 300,
}
if 'postgresql' in app.config['SQLALCHEMY_DATABASE_URI']:
    engine_options['connect_args'] = {
        'connect_timeout': 10,
        'sslmode': 'require'
    }
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = engine_options

# Fix for Render PostgreSQL URLs
if app.config['SQLALCHEMY_DATABASE_URI'].startswith('postgres://'):
    app.config['SQLALCHEMY_DATABASE_URI'] = app.config['SQLALCHEMY_DATABASE_URI'].replace('postgres://', 'postgresql://', 1)

# Google OAuth Configuration
app.config['GOOGLE_CLIENT_ID'] = os.getenv('GOOGLE_CLIENT_ID')
app.config['GOOGLE_CLIENT_SECRET'] = os.getenv('GOOGLE_CLIENT_SECRET')

# Email Configuration
app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER', 'smtp.gmail.com')
app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', 587))
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
app.config['CONTACT_EMAIL'] = 'francisj2nd@gmail.com'

# Load environment variables from Render
GCP_LOCATION = os.getenv("GCP_LOCATION", "us-central1")
GCP_PROJECT_ID = os.getenv("GCP_PROJECT_ID")
MODEL_NAME = os.getenv("MODEL_NAME", "gemini-2.5-flash-lite")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GOOGLE_CSE_ID = os.getenv("GOOGLE_CSE_ID")

# Usage limits for free plan
MONTHLY_WORD_LIMIT = 15000
MONTHLY_DOWNLOAD_LIMIT = 10

# Initialize extensions
db.init_app(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'auth_login'
login_manager.login_message = 'Please log in to access this page.'
login_manager.login_message_category = 'info'

@login_manager.user_loader
def load_user(user_id):
    try:
        return User.query.get(int(user_id))
    except Exception as e:
        logger.error(f"Error loading user {user_id}: {e}")
        return None

# Database connection retry decorator
def retry_db_operation(max_retries=3, delay=1):
    def decorator(func):
        def wrapper(*args, **kwargs):
            for attempt in range(max_retries):
                try:
                    return func(*args, **kwargs)
                except (OperationalError, DatabaseError) as e:
                    logger.warning(f"Database operation failed (attempt {attempt + 1}/{max_retries}): {e}")
                    if attempt < max_retries - 1:
                        time.sleep(delay * (2 ** attempt))  # Exponential backoff
                        continue
                    else:
                        logger.error(f"Database operation failed after {max_retries} attempts: {e}")
                        raise
                except Exception as e:
                    logger.error(f"Unexpected error in database operation: {e}")
                    raise
            return None
        return wrapper
    return decorator

# Validation check - log missing variables but don't fail
missing_vars = []
if not GCP_PROJECT_ID: missing_vars.append("GCP_PROJECT_ID")
if not GOOGLE_API_KEY: missing_vars.append("GOOGLE_API_KEY")
if not GOOGLE_CSE_ID: missing_vars.append("GOOGLE_CSE_ID")
if not os.getenv("GOOGLE_APPLICATION_CREDENTIALS"): missing_vars.append("GOOGLE_APPLICATION_CREDENTIALS")

if missing_vars:
    logger.warning(f"Missing environment variables: {', '.join(missing_vars)}")

# Studio Type Mapping
STUDIO_TYPE_MAPPING = {
    'article': ['ARTICLE'],
    'social': ['SOCIAL_POST', 'EMAIL', 'AD_COPY'],
    'editing': ['TEXT_REFINEMENT', 'SUMMARY', 'TRANSLATION'],
    'repurpose': ['REPURPOSE_TWEET', 'REPURPOSE_SLIDES'],
    'seo': ['SEO_KEYWORDS', 'SEO_HEADLINES'],
    'business': ['PRESS_RELEASE', 'JOB_DESCRIPTION'],
    'brainstorming': ['IDEAS'],
    'scriptwriting': ['SCRIPT'],
    'ecommerce': ['ECOMMERCE'],
    'webcopy': ['WEBCOPY'],
}

# --- Initialize genai client with Vertex AI ---
CLIENT = None
try:
    if GCP_PROJECT_ID and os.getenv("GOOGLE_APPLICATION_CREDENTIALS"):
        vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
        CLIENT = GenerativeModel(model_name=MODEL_NAME)
        logger.info("Google AI Client Initialized Successfully via Vertex AI.")
    else:
        logger.warning("Google AI client not initialized - missing credentials")
except Exception as e:
    logger.error(f"Failed to initialize Google AI client: {e}")

# Helper functions
def construct_initial_prompt(topic, settings=None):
    if settings is None:
        settings = {}

    # Start building the prompt
    prompt_lines = [
        f"I want you to generate a high-quality, SEO-optimized thought-leadership article on the topic of: \"{topic}\"",
        "\n**Article Requirements:**"
    ]

    # Dynamically add instructions based on settings
    word_count = settings.get('wordCount', '1000')
    prompt_lines.append(f"1. **Length:** Aim for approximately {word_count} words.")

    tone = settings.get('tone', 'Professional')
    prompt_lines.append(f"2. **Tone of Voice:** The article's tone must be {tone}.")

    audience = settings.get('audience')
    if audience:
        prompt_lines.append(f"3. **Target Audience:** Write for an audience of {audience}.")

    prompt_lines.extend([
        "4. **Structure:** Use clear sections with H2 and H3 subheadings using Markdown.",
        "5. **Placeholders:** Include 3 relevant image placeholders. For each, provide a suggested title and a full, SEO-optimized alt text. Format them exactly like this: `[Image Placeholder: Title, Alt Text]`",
        "6. **Quality:** The content must be original, human-readable, and valuable."
    ])

    key_points = settings.get('keyPoints')
    if key_points:
        prompt_lines.append("\n**Key Points to Include:**")
        prompt_lines.append("The article must incorporate and expand upon the following key points:")
        prompt_lines.append(key_points)

    seo_keyword = settings.get('seoKeyword')
    if seo_keyword:
        prompt_lines.append("\n**SEO Optimization:**")
        prompt_lines.append(f"The primary SEO keyword to focus on is \"{seo_keyword}\". Please integrate this keyword naturally throughout the article, including in headings where appropriate.")

    cta = settings.get('cta')
    if cta:
        prompt_lines.append("\n**Call to Action:**")
        prompt_lines.append(f"Conclude the article with the following call to action: \"{cta}\"")

    # Add the standard closing for SEO elements
    prompt_lines.append("\n**Final Output:**")
    prompt_lines.append("At the very end of the article, after all other content, provide \"SEO Keywords:\" and \"Meta Description:\".")

    return "\n".join(prompt_lines)

def construct_social_post_prompt(topic, goal, platform, settings=None):
    """Constructs a prompt for generating social media posts."""
    if settings is None:
        settings = {}

    variations = settings.get('variations', 3)
    tone = settings.get('tone', 'professional')
    emojis = "Include relevant emojis." if settings.get('emojis', True) else "Do not use any emojis."
    cta = settings.get('cta')

    prompt = f"""
You are a social media marketing expert. Your task is to generate {variations} distinct social media posts for the {platform} platform.

**Topic:** {topic}
**Goal of the posts:** {goal}
**Tone:** {tone}

**Requirements:**
1.  **Platform:** Tailor the tone and length for {platform}. For Twitter, keep it concise. For LinkedIn, be more professional.
2.  **Variety:** Provide {variations} distinct options with different angles or hooks.
3.  **Hashtags:** Include relevant and trending hashtags for each post.
4.  **Emojis:** {emojis}
"""
    if cta:
        prompt += f"5.  **Call to Action/Link:** Include the following call to action or link: {cta}\n"

    prompt += "6.  **Formatting:** Present each post clearly separated. Use \"---\" on a new line between each post option.\n\nGenerate the social media posts now."
    return prompt

def construct_email_prompt(topic, audience, tone, settings=None):
    """Constructs a prompt for generating an email."""
    if settings is None:
        settings = {}

    email_type = settings.get('emailType', 'Promotional')
    brand_name = settings.get('brandName')
    cta_button = settings.get('ctaButton')

    prompt = f"""
You are an expert copywriter specializing in email marketing. Your task is to write a compelling {email_type} email.

**Topic/Product:** {topic}
**Target Audience:** {audience}
**Desired Tone:** {tone}
"""
    if brand_name:
        prompt += f"**Brand Name:** {brand_name}\n"

    prompt += """
**Requirements:**
1.  **Subject Line:** Generate 3-5 engaging subject line options.
2.  **Email Body:** Write a complete email body based on the topic. It should have a clear hook, body, and a strong call-to-action (CTA).
"""
    if cta_button:
        prompt += f"3.  **CTA Button:** The final call to action should ideally lead to a button with the text: \"{cta_button}\".\n"

    prompt += "4.  **Formatting:** Use Markdown for formatting (e.g., bolding, bullet points). Start with the subject lines, followed by \"---\", then the email body.\n\nGenerate the email content now."
    return prompt

def construct_refine_text_prompt(text_to_refine, settings=None):
    """Constructs a prompt for refining text to a specific tone."""
    if settings is None:
        settings = {}

    goal = settings.get('goal', 'Improve Clarity')
    formality = settings.get('formality', 'Neutral')
    audience = settings.get('audience')

    prompt = f"""
You are an expert editor. Your task is to revise the following text.

**Original Text:**
---
{text_to_refine}
---

**Instructions:**
1.  **Primary Goal:** Your main goal is to "{goal}".
2.  **Formality:** The revised text should have a "{formality}" formality level.
"""
    if audience:
        prompt += f"3.  **Target Audience:** Adapt the text to be easily understood by an audience of {audience}.\n"

    prompt += "4.  **Core Meaning:** Do not add any new information or change the core meaning of the text.\n"
    prompt += "5.  **Output:** Return only the revised text, without any commentary or preamble.\n\nRevise the text now."
    return prompt

def construct_repurpose_prompt(repurpose_type, original_content, settings=None):
    """Constructs a dynamic prompt for various content repurposing tasks."""
    if settings is None:
        settings = {}

    # Universal settings
    audience = settings.get('audience', 'a general audience')
    goal = settings.get('goal', 'to inform and engage')
    message = settings.get('message')

    # Base prompt
    prompt = f"""
You are a strategic content repurposing expert. Your task is to transform the following original content into a new format.

**Original Content:**
---
{original_content}
---

**Repurposing Goal:**
-   **New Format:** {repurpose_type.replace('_', ' ').title()}
-   **Target Audience:** {audience}
-   **Primary Goal:** {goal}
"""
    if message:
        prompt += f"-   **Key Message to Emphasize:** {message}\n"

    prompt += "\n**Format-Specific Instructions:**\n"

    # Dynamic instructions based on type
    if repurpose_type == 'twitter_thread':
        tweet_count = settings.get('tweetCount', 5)
        hook = "**Include a strong, engaging hook** in the first tweet." if settings.get('includeHook', True) else ""
        hashtags = "**Include relevant hashtags** in the final tweet." if settings.get('includeHashtags', True) else ""
        prompt += f"""
-   **Structure:** Create a numbered Twitter thread with {tweet_count} tweets.
-   **Content:** Summarize the key points of the original content. {hook} {hashtags}
-   **Constraints:** Each tweet must be under 280 characters.
-   **Separator:** Use "---" on a new line to separate each tweet.
"""
    elif repurpose_type == 'linkedin_post':
        professional_tone = "Maintain a professional and business-oriented tone." if settings.get('useProfessionalTone', True) else ""
        prompt += f"""
-   **Structure:** A single, well-formatted LinkedIn post.
-   **Content:** Summarize the key takeaways for a professional audience. {professional_tone}
-   **Formatting:** Use bullet points or numbered lists for readability. Include 3-5 relevant hashtags.
"""
    elif repurpose_type == 'video_script':
        target_length = settings.get('targetLength', '3-5 minutes')
        spoken_by = settings.get('spokenBy', 'a single narrator')
        visual_cues = "**Include suggested visual cues** or B-roll shots in parentheses where appropriate (e.g., `(Show a chart of Q3 growth)`)." if settings.get('includeVisualCues', True) else ""
        prompt += f"""
-   **Structure:** A video script with clear sections (Hook, Intro, Main Points, Outro).
-   **Target Length:** The script should be paced for a video of approximately {target_length}.
-   **Dialogue:** The script should be written for {spoken_by}.
-   **Content:** {visual_cues}
"""

    prompt += "\n**Final Output:**\nReturn only the repurposed content as requested, without any of your own commentary or preamble."
    return prompt

def construct_keyword_strategy_prompt(settings=None):
    """Constructs a prompt for a comprehensive keyword strategy."""
    if settings is None:
        settings = {}

    description = settings.get('description')
    audience = settings.get('audience')
    intents = ", ".join(settings.get('intents', ['Informational']))
    competitors = settings.get('competitors')

    prompt = f"""
You are a world-class SEO strategist. Your task is to develop a comprehensive keyword strategy based on the following business context.

**Business Context:**
-   **Description:** {description}
-   **Target Audience:** {audience}
-   **Desired Search Intent Focus:** {intents}

**Instructions:**
1.  **Core Keyword Clusters:** Identify 3-5 primary keyword clusters (themes) central to the business.
2.  **Keyword Generation:** For each cluster, generate a list of relevant keywords, including:
    *   `primary_keywords`: (2-3 words) High-volume, core terms.
    *   `long_tail_keywords`: (4+ words) More specific, lower-volume phrases.
    *   `question_keywords`: Common questions the target audience asks.
"""
    if competitors:
        prompt += f"""
3.  **Competitive Gap Analysis:**
    *   **Analyze Competitors:** Briefly analyze the keyword strategy of the following competitors: {competitors}
    *   **Identify Opportunities:** Suggest 5-10 "gap" keywords that the competitors are ranking for but are also attainable and relevant for our business.
"""
    prompt += """
4.  **Format as JSON:** Return the entire strategy as a single, valid JSON object. The main keys should be "keyword_clusters" and, if applicable, "competitive_gap_analysis".
5.  **Return Only JSON:** Do not include any preamble, commentary, or markdown formatting.

Generate the keyword strategy now.
"""
    return prompt

def construct_seo_audit_prompt(settings=None):
    """Constructs a prompt for an on-page SEO audit."""
    if settings is None:
        settings = {}

    url = settings.get('url')
    primary_keyword = settings.get('primaryKeyword')
    secondary_keywords = settings.get('secondaryKeywords')
    page_goal = settings.get('pageGoal')
    competitor_url = settings.get('competitorUrl')

    prompt = f"""
You are a senior technical SEO analyst. Your task is to perform a detailed on-page SEO audit for the given URL, focusing on outranking competitors.

**Audit Target:**
-   **URL:** {url}
-   **Primary Target Keyword:** "{primary_keyword}"
"""
    if secondary_keywords:
        prompt += f"-   **Secondary Keywords:** {secondary_keywords}\n"
    if page_goal:
        prompt += f"-   **Desired Page Goal/CTA:** {page_goal}\n"
    if competitor_url:
        prompt += f"-   **Primary Competitor to Outrank:** {competitor_url}\n"

    prompt += """
**Instructions:**
Provide a step-by-step audit in Markdown format. For each point, provide a "Current State" (if you can infer it), an "Assessment" (Good, Needs Improvement, Poor), and a specific, actionable "Recommendation".

**Audit Checklist:**
1.  **Title Tag:**
2.  **Meta Description:**
3.  **H1 Tag:**
4.  **Subheadings (H2, H3):**
5.  **Keyword Usage:** (Primary and secondary keywords)
6.  **Content Quality & Depth:**
7.  **URL Structure:**
8.  **Internal Linking:**
"""
    if competitor_url:
        prompt += """
9.  **Competitive Analysis & Recommendations:**
    *   Analyze the on-page SEO of the competitor's URL.
    *   Provide 3-5 specific, actionable recommendations on how our page can be improved to outperform them for the target keywords.
"""
    prompt += "\nReturn only the audit in Markdown format, without any preamble."
    return prompt

def construct_brainstorm_prompt(settings=None):
    """Constructs a prompt for general brainstorming."""
    if settings is None:
        settings = {}

    goal = settings.get('goal')
    framework = settings.get('framework', 'Simple List')
    constraints = settings.get('constraints')

    prompt = f"""
You are a strategic facilitator. Your task is to lead a brainstorming session based on the following objective and framework.

**Objective:** {goal}
**Framework:** {framework}
"""
    if constraints:
        prompt += f"**Constraints:** {constraints}\n"

    prompt += """
**Instructions:**
-   Generate a list of creative, relevant, and actionable ideas.
-   If using SWOT, structure the output with 'Strengths', 'Weaknesses', 'Opportunities', and 'Threats' headings.
-   If using a Simple List, provide a direct list of ideas.
-   Return only the generated ideas, without any preamble.
"""
    return prompt

def construct_naming_prompt(settings=None):
    """Constructs a prompt for generating names."""
    if settings is None:
        settings = {}

    description = settings.get('description')
    style = settings.get('style', 'Descriptive')
    tone = settings.get('tone', 'Modern')
    keyword = settings.get('keyword')
    length = settings.get('length', 'Any')

    prompt = f"""
You are an expert branding and naming consultant. Your task is to generate a list of 10-15 potential names for a product or company.

**Product/Company Description:** {description}

**Naming Guidelines:**
-   **Style:** {style}
-   **Tone:** {tone}
-   **Length/Syllables:** {length}
"""
    if keyword:
        prompt += f"-   **Required Keyword:** Must include the word or concept '{keyword}'.\n"

    prompt += """
**Instructions:**
-   Provide a diverse list of names that fit the guidelines.
-   The names should be memorable, easy to spell, and ideally have domain name potential.
-   Return only the list of names, each on a new line, without any commentary.
"""
    return prompt

def construct_script_prompt(topic, settings=None):
    """Constructs a dynamic prompt for generating a script based on detailed user settings."""
    if settings is None:
        settings = {}

    # Extract settings with defaults
    format_map = {
        "youtube_video": "YouTube Video",
        "tiktok_shorts": "TikTok/Shorts Video",
        "podcast_episode": "Podcast Episode"
    }
    script_format = format_map.get(settings.get('format'), "YouTube Video")
    audience = settings.get('audience')
    tone = settings.get('tone')
    characters = settings.get('characters')
    elements = settings.get('structural_elements', [])

    # Start building the prompt
    prompt = f"""
You are a professional scriptwriter. Your task is to create a detailed, engaging script for a "{script_format}" based on the user's specifications.

**Core Topic/Premise:** {topic}
"""
    # Add optional context
    if audience:
        prompt += f"**Target Audience:** {audience}\n"
    if tone:
        prompt += f"**Desired Tone/Style:** {tone}\n"
    if characters:
        # Sanitize the input to remove potentially problematic newline characters
        sane_characters = characters.replace('\n', ', ')
        prompt += f"**Characters/Hosts:**\n- {sane_characters}\n"

    # Add structural requirements
    prompt += "\n**Script Requirements:**\n"
    prompt += f"1.  **Format:** The script must be written specifically for a **{script_format}**. This means short, punchy segments for TikTok, a classic structure for YouTube, or a conversational flow for a podcast.\n"
    prompt += "2.  **Structure:** The script must be well-organized. Use Markdown for headings (e.g., `## Scene 1`, `### Alex's Monologue`).\n"

    # Add character dialogue instructions
    if characters:
        prompt += "3.  **Dialogue:** Clearly denote which character is speaking (e.g., `ALEX: [dialogue]`).\n"

    # Add instructions for each selected structural element
    element_instructions = {
        "hook": "Include a powerful, attention-grabbing hook at the very beginning.",
        "cta": "Include a clear Call to Action at the end of the script (e.g., asking to subscribe, comment, or visit a website).",
        "ad_break": "Include a placeholder for an ad break, like `[AD BREAK - 60 seconds]`.",
        "visual_cues": "If this is a visual format (YouTube/TikTok), include suggested visual cues, camera angles, or on-screen text in parentheses, like `(Close up on the product)`.",
        "sfx": "Include suggestions for sound effects where they would enhance the script, like `(Sound of a dramatic whoosh)`."
    }

    for i, element_key in enumerate(elements, start=4):
        if element_key in element_instructions:
            prompt += f"{i}.  **{element_key.replace('_', ' ').title()}:** {element_instructions[element_key]}\n"

    prompt += "\nGenerate the script now, adhering to all the requirements above. Return only the script content."
    return prompt

def construct_product_description_prompt(settings=None):
    """Constructs a prompt for generating a product description."""
    if settings is None: settings = {}

    prompt = f"""
You are an expert e-commerce copywriter with a knack for persuasive, benefit-driven language.
Your task is to write a compelling product description for the following product.

**Product Name:** {settings.get('productName', '')}
**Target Audience:** {settings.get('audience', '')}
**Tone of Voice:** {settings.get('tone', '')}
**Primary Benefit to Emphasize:** {settings.get('benefit', '')}

**Key Features:**
---
{settings.get('features', '')}
---

**Instructions:**
1.  **Analyze the Features:** Do not just list the features. For each one, explain the *benefit* it provides to the target audience.
2.  **Focus on the Primary Benefit:** Ensure the primary benefit, "{settings.get('benefit', '')}", is the central theme of the description.
3.  **Adopt the Tone:** The entire description must be written in a {settings.get('tone', '')} tone.
4.  **Structure:** Format the output as a "{settings.get('format', 'paragraph')}". If using bullets, make them benefit-oriented.
5.  **Output:** Return only the final product description, without any of your own commentary.

Generate the product description now.
"""
    return prompt

def construct_campaign_prompt(settings=None):
    """Constructs a prompt for generating a promotional campaign."""
    if settings is None: settings = {}

    assets_text = ", ".join(settings.get('assets', ['email', 'social']))

    prompt = f"""
You are a senior marketing manager tasked with creating a new promotional campaign.

**Campaign Details:**
- **Product/Service:** {settings.get('product', '')}
- **Occasion:** {settings.get('occasion', '')}
- **The Offer:** {settings.get('offer', '')}
- **Urgency Element:** {settings.get('urgency', '')}

**Task:**
Generate a cohesive set of marketing assets for this campaign. The required assets are: **{assets_text}**.

**Instructions:**
1.  **Create a Section for Each Asset:** Use a clear Markdown heading (e.g., `### Email Copy`) for each requested asset.
2.  **Maintain Cohesion:** Ensure the messaging, tone, and offer are consistent across all assets.
3.  **Incorporate Urgency:** Weave the urgency element ("{settings.get('urgency', '')}") into the copy where appropriate to drive action.
4.  **Action-Oriented:** All copy should be persuasive and guide the customer towards making a purchase with the offer.
5.  **Output:** Return only the generated marketing assets, without any of your own commentary.

Generate the campaign assets now.
"""
    return prompt

def construct_review_response_prompt(settings=None):
    """Constructs a prompt for responding to a customer review."""
    if settings is None: settings = {}

    prompt = f"""
You are a senior customer support manager. Your goal is to respond to a customer review in a way that is helpful, on-brand, and resolves any issues.

**Customer Review Details:**
- **Star Rating Given:** {settings.get('rating', '3')}/5
- **Review Content:** "{settings.get('review', '')}"

**Response Requirements:**
- **Tone:** Your response must be **{settings.get('tone', 'Empathetic & Helpful')}**.
- **Action/Offer:** {settings.get('offer') if settings.get('offer') else 'No special offer is required.'}

**Instructions:**
1.  **Analyze Sentiment:** Based on the star rating and the review content, correctly identify the customer's sentiment (e.g., happy, frustrated, confused).
2.  **Address Key Points:** Directly address the specific points, positive or negative, mentioned in the review.
3.  **Incorporate the Tone:** Maintain the required tone throughout the entire response.
4.  **Integrate the Offer:** If an action or offer is specified (e.g., "Offer a 15% discount"), seamlessly weave it into the response as a gesture of goodwill or a solution. Do not make it sound like a bribe.
5.  **Output:** Return only the customer-facing response, without any of your own commentary or analysis.

Generate the customer review response now.
"""
    return prompt

def construct_ecommerce_prompt(name, features, tone):
    """Constructs a prompt for generating a product description."""
    return f"""
    You are an expert e-commerce copywriter. Your task is to write a persuasive, benefit-focused product description.

    **Product Name:** {name}

    **Key Features:**
    {features}

    **Target Tone:** {tone}

    **Instructions:**
    1.  **Create a Compelling Narrative:** Start with a hook that grabs attention.
    2.  **Translate Features to Benefits:** For each feature, explain how it benefits the customer. Don't just list features; sell the solution and experience.
    3.  **Incorporate Tone:** Write the entire description in a {tone} tone.
    4.  **Structure:** Use a short introduction, a few paragraphs or bullet points for the benefits, and a concluding sentence that encourages a purchase.
    5.  **Return Only the Description:** Do not include any preamble or commentary.

    Generate the product description now.
    """

def construct_landing_page_prompt(settings=None):
    """Constructs a prompt for generating a full landing page."""
    if settings is None: settings = {}

    prompt = f"""
You are an expert direct-response copywriter. Your task is to write a complete set of copy for a landing page with a single, clear goal.

**Landing Page Details:**
- **Product Name:** {settings.get('productName', '')}
- **Target Audience:** {settings.get('audience', '')}
- **Page Goal:** {settings.get('goal', '')}
- **Primary Pain Point to Solve:** {settings.get('painPoint', '')}
- **Tone:** {settings.get('tone', '')}

**Key Features:**
---
{settings.get('features', '')}
---

**Instructions:**
1.  **Solve the Pain Point:** All copy must focus on how the product's features solve the primary pain point for the target audience.
2.  **Benefit-Oriented:** Do not just list features; explain the tangible benefits and outcomes for the user.
3.  **Clear Structure:** Structure the output using the following Markdown headings: `### Headline`, `### Sub-headline`, `### Body`, and `### Call-to-Action`.
4.  **Persuasive Body:** The Body copy should be compelling, persuasive, and directly encourage the user to achieve the page goal.
5.  **Actionable CTA:** The Call-to-Action should be a short, punchy phrase that tells the user exactly what to do (e.g., "Get Started Free," "Claim Your Discount").
6.  **Output:** Return only the structured copy, without any of your own commentary.

Generate the landing page copy now.
"""
    return prompt

def construct_homepage_section_prompt(settings=None):
    """Constructs a dynamic prompt for generating a specific homepage section."""
    if settings is None: settings = {}

    section_type = settings.get('sectionType', 'hero')
    section_map = {
        "hero": "Hero Section",
        "features": "Features/Benefits Section",
        "how_it_works": "How It Works Section",
        "faq": "FAQ Section"
    }
    section_name = section_map.get(section_type, "Hero Section")

    prompt = f"""
You are a specialist web copywriter. Your task is to write the copy for a specific section of a company's homepage.

**Company Name:** {settings.get('companyName', '')}
**Company One-Liner:** {settings.get('oneLiner', '')}
**Section to Generate:** {section_name}

**Key Information Provided by User:**
---
{settings.get('keyInfo', '')}
---

**Instructions:**
Your output must be structured correctly for the requested section.
"""

    if section_type == 'hero':
        prompt += """
- **Task:** Write a powerful, attention-grabbing headline, a clear and concise sub-headline that expands on the headline, and two distinct, compelling CTA button texts.
- **Structure:** Use the following Markdown headings for each part: `#### Headline`, `#### Sub-headline`, `#### Primary CTA`, `#### Secondary CTA`.
"""
    elif section_type == 'features':
        prompt += """
- **Task:** Write a main headline for the features section. Then, for 3-5 key features, write a short, benefit-focused title and a 1-2 sentence description for each.
- **Structure:** Use a main `### Features Section Headline`. For each feature, use a `#### Feature Title` and a paragraph for the description.
"""
    elif section_type == 'how_it_works':
        prompt += """
- **Task:** Write a main headline for this section. Then, write copy for 3-4 simple steps that explain how the service works.
- **Structure:** Use a main `### How It Works Headline`. For each step, use a `#### Step X: [Step Title]` heading followed by a short descriptive paragraph.
"""
    elif section_type == 'faq':
        prompt += """
- **Task:** Write a main headline for the FAQ section. Then, based on the user's key information, write 4-5 questions and their corresponding clear, concise answers.
- **Structure:** Use a main `### FAQ Section Headline`. For each question, use a `#### Question:` heading and a paragraph for the answer.
"""

    prompt += "\nReturn only the generated copy in the requested structure, without any of your own commentary."
    return prompt

def construct_usp_prompt(settings=None):
    """Constructs a prompt for generating a Value Proposition and USP."""
    if settings is None: settings = {}

    prompt = f"""
You are a senior brand strategist. Your task is to analyze the following product information and generate a powerful Value Proposition and a list of Unique Selling Propositions (USPs).

**Product Information:**
- **Product Description:** {settings.get('productDesc', '')}
- **Target Customer:** {settings.get('customer', '')}
- **Key Differentiators:** {settings.get('differentiators', '')}
- **Competitors:** {settings.get('competitors', 'Not specified')}

**Instructions:**
1.  **Analyze the Data:** Carefully consider how the product's features and differentiators appeal to the target customer, especially in contrast to the competitors.
2.  **Generate Value Proposition:** Create a single, clear, and compelling sentence that explains the core benefit the product provides to the customer. This should be customer-centric and outcome-focused.
3.  **Generate USPs:** Create a list of 3-5 concise, memorable, and powerful statements that highlight the key differentiators. These are the specific, unique reasons a customer should choose this product over others.
4.  **Structure the Output:** Use the following Markdown headings: `### Value Proposition` and `### Unique Selling Propositions (USPs)`. List the USPs as a bulleted list.
5.  **Output:** Return only the generated propositions, without any of your own commentary.

Generate the Value Proposition and USPs now.
"""
    return prompt

def construct_proposal_prompt(settings=None):
    """Constructs a prompt for generating a business proposal."""
    if settings is None: settings = {}
    prompt = f"""
You are a professional business consultant and proposal writer. Your task is to generate a formal and persuasive business proposal based on the provided details.

**Proposal Details:**
- **Your Company:** {settings.get('company', '')}
- **Client Name:** {settings.get('client', '')}
- **Client's Problem:** {settings.get('problem', '')}
- **Proposed Solution:** {settings.get('solution', '')}
- **Key Deliverables:**
{settings.get('deliverables', '')}
- **Desired Tone:** {settings.get('tone', 'Formal and confident')}

**Instructions:**
1.  **Structure the Proposal:** Generate a complete proposal with the following sections, using clear Markdown headings:
    - `## Introduction`: Briefly introduce your company and the purpose of the proposal.
    - `## Understanding the Problem`: Show you understand the client's needs by summarizing their problem.
    - `## Proposed Solution`: Detail your proposed solution, explaining how it addresses the client's problem.
    - `## Deliverables & Timeline`: List the key deliverables and provide a high-level timeline.
    - `## Conclusion`: End with a confident closing statement and a call to action (e.g., "We look forward to discussing this proposal with you further.").
2.  **Adopt the Tone:** The entire proposal must be written in a **{settings.get('tone', 'Formal and confident')}** tone.
3.  **Output:** Return only the complete proposal text. Do not include any of your own commentary.

Generate the business proposal now.
"""
    return prompt

def construct_report_prompt(settings=None):
    """Constructs a dynamic prompt for generating a formal report."""
    if settings is None: settings = {}

    report_type = settings.get('reportType', 'status_update')
    report_name = report_type.replace('_', ' ').title()

    prompt = f"""
You are a senior business analyst. Your task is to write a clear, structured, and formal **{report_name}**.

**Report Details:**
- **Subject:** {settings.get('subject', '')}
- **Time Period:** {settings.get('period', '')}
- **Target Audience:** {settings.get('audience', '')}

**Key Data Points / Information to Include:**
---
{settings.get('dataPoints', '')}
---

**Instructions:**
1.  **Adopt the Persona:** Write from the perspective of a professional analyst reporting to the specified audience.
2.  **Convert Data to Narrative:** Do not just list the data points. Weave them into a professional, easy-to-understand narrative.
3.  **Use a Logical Structure:** The report must be well-structured. Use the following structure based on the report type:
"""

    if report_type == 'status_update':
        prompt += """
    - `## 1. Executive Summary`: A brief overview of the project status.
    - `## 2. Accomplishments`: Detail what has been achieved during this period, referencing the data points.
    - `## 3. Challenges & Roadblocks`: Outline any issues encountered.
    - `## 4. Next Steps`: Describe the planned activities for the next period.
"""
    elif report_type == 'financial_summary':
        prompt += """
    - `## 1. Overview`: A high-level summary of the financial performance for the period.
    - `## 2. Key Metrics`: Present the key data points with brief explanations.
    - `## 3. Analysis & Insights`: Provide analysis on what the data means (e.g., trends, anomalies).
    - `## 4. Outlook`: Provide a brief forecast or outlook based on the summary.
"""
    elif report_type == 'incident_report':
        prompt += """
    - `## 1. Summary of Incident`: What happened, when, and where.
    - `## 2. Impact Assessment`: Detail the impact of the incident, using the provided data.
    - `## 3. Root Cause Analysis`: Explain what caused the incident.
    - `## 4. Resolution and Next Steps`: Describe the actions taken to resolve the issue and prevent recurrence.
"""

    prompt += "\n4. **Output:** Return only the complete, structured report. Do not include any of your own commentary."
    return prompt

def construct_press_release_prompt(settings=None):
    """Constructs a prompt for generating a professional press release."""
    if settings is None: settings = {}

    prompt = f"""
You are a public relations (PR) professional. Your task is to write a professional press release in strict AP style based on the provided information.

**Press Release Details:**
- **Headline:** {settings.get('headline', '')}
- **Company Info (for dateline):** {settings.get('companyInfo', '')}
- **Key Information (5 Ws):** {settings.get('keyInfo', '')}
- **Quote (with speaker and title):** {settings.get('quote', '')}
- **Company Boilerplate:** {settings.get('boilerplate', '')}

**Instructions:**
1.  **Strict AP Format:** The output MUST follow the standard Associated Press (AP) style for a press release.
2.  **Structure:** The document must include the following elements in this exact order:
    - `FOR IMMEDIATE RELEASE` (all caps).
    - The **Headline** you are given.
    - A **Dateline** in the format `CITY, State – (Date) –`.
    - An **Introduction (Lede):** The first paragraph must summarize the most critical information from the 5 Ws.
    - A **Body:** Subsequent paragraphs that elaborate on the key information and seamlessly integrate the provided **Quote**.
    - An **About/Boilerplate Section:** A paragraph starting with "About [Company Name]" using the provided boilerplate text.
    - **Media Contact:** A placeholder for media contact information (e.g., `Media Contact: [Name] [Email]`).
    - A **Sign-off:** The document must end with three hash symbols (`###`) on a new line.
3.  **Output:** Return only the complete, formatted press release. Do not include any of your own commentary.

Generate the press release now.
"""
    return prompt

def construct_ad_copy_prompt(product, audience, settings=None):
    """Constructs a prompt for generating ad copy."""
    if settings is None:
        settings = {}

    platform = settings.get('platform', 'Google Ads')
    key_benefit = settings.get('keyBenefit')
    tone = settings.get('tone', 'Urgent')

    platform_constraints = {
        'Google Ads': 'Headline max 30 chars, Description max 90 chars.',
        'Facebook/Instagram': 'Keep copy visual-focused and engaging. Use emojis.',
        'LinkedIn Ads': 'Maintain a professional and business-oriented tone.'
    }
    constraint = platform_constraints.get(platform, '')

    prompt = f"""
You are a senior copywriter at a top advertising agency. Your task is to generate 3 distinct pieces of ad copy for a given product and audience, optimized for the {platform} platform.

**Product/Service:** {product}
**Target Audience:** {audience}
**Desired Tone:** {tone}
"""
    if key_benefit:
        prompt += f"**Key Benefit/Pain Point to Address:** {key_benefit}\n"

    prompt += f"""
**Instructions:**
1.  **Generate 3 Variations:** Create three distinct ad copy options. Each should have a different angle or hook.
2.  **Structure:** For each variation, provide a "Headline:" and a "Body:".
3.  **Platform Constraints:** Adhere to the following constraints for {platform}: {constraint}
4.  **Clarity and Persuasion:** The copy must be clear, concise, and persuasive.
5.  **Separator:** Use "---" on a new line to separate each ad copy variation.
6.  **Return Only the Ad Copy:** Do not include any preamble or commentary.

Generate the ad copy now.
"""
    return prompt

def construct_summarizer_prompt(text, settings=None):
    """Constructs a prompt for summarizing text."""
    if settings is None:
        settings = {}

    summary_format = settings.get('format', 'Paragraph')
    extraction_type = settings.get('extractionType', 'Summarize')
    focus_area = settings.get('focus')

    prompt = f"""
You are an expert at analyzing and summarizing text. Your task is to process the following text according to the user's specific instructions.

**Original Text:**
---
{text}
---

**Task Details:**
1.  **Extraction Type:** Your primary task is to "{extraction_type}".
2.  **Output Format:** Present the result as a "{summary_format}".
"""
    if focus_area:
        prompt += f"3.  **Focus Area:** Pay special attention to and prioritize information related to: \"{focus_area}\".\n"

    prompt += """
**Instructions:**
-   If summarizing, distill the core ideas and most important information.
-   If extracting action items, list clear, actionable tasks.
-   If extracting statistics, list all key data points and numbers.
-   The output must be clear, accurate, and easy to understand.
-   Return only the requested output, without any preamble or commentary.

Generate the result now.
"""
    return prompt

def construct_translator_prompt(text, settings=None):
    """Constructs a prompt for translating and localizing text."""
    if settings is None:
        settings = {}

    locale = settings.get('locale', 'Spanish (Mexico)')
    formality = settings.get('formality', 'Formal')
    adapt_idioms = settings.get('adaptIdioms', True)

    idiom_instruction = (
        "It is crucial that you adapt idioms, cultural references, and colloquialisms to be natural-sounding and culturally appropriate for the target locale. Do not translate them literally."
        if adapt_idioms
        else "Translate the text as literally as possible, preserving original idioms and phrasing, even if they don't sound natural in the target language."
    )

    prompt = f"""
You are an expert localization specialist, not just a literal translator. Your task is to translate and localize the following text with cultural and contextual accuracy.

**Text to Localize:**
---
{text}
---

**Localization Task:**
1.  **Target Locale:** {locale}
2.  **Formality Level:** {formality}
3.  **Cultural Adaptation:** {idiom_instruction}

**Instructions:**
-   Ensure the final text is grammatically correct and fluent for a native speaker of the {locale} locale.
-   Maintain the core meaning and intent of the original text.
-   Return only the final, localized text, without any of your own commentary, preambles, or the original text.

Localize the text now.
"""
    return prompt


def get_image_url(query):
    if not GOOGLE_API_KEY or not GOOGLE_CSE_ID:
        logger.warning("Google Search API credentials not configured.")
        return None

    api_url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "key": GOOGLE_API_KEY,
        "cx": GOOGLE_CSE_ID,
        "q": query,
        "searchType": "image",
        "num": 1
    }
    try:
        response = requests.get(api_url, params=params, timeout=5)
        response.raise_for_status()
        data = response.json()
        if "items" in data and len(data["items"]) > 0:
            item = data["items"][0]
            return {"url": item["link"], "source": item["image"]["contextLink"]}
        return None
    except requests.RequestException as e:
        logger.error(f"Error fetching image from Google: {e}")
        return None

def format_article_content(raw_markdown_text, topic="", enable_images=False):
    """Format article content with images"""
    hybrid_content = raw_markdown_text
    placeholder_regex = re.compile(r"\[Image Placeholder: (.*?),\s*(.*?)\]")

    if enable_images:
        for match in placeholder_regex.finditer(raw_markdown_text):
            original_placeholder = match.group(0)
            title = match.group(1).strip()
            alt_text = match.group(2).strip()
            image_data = get_image_url(alt_text)

            if image_data:
                new_image_tag = (
                    f'<div class="real-image-container">'
                    f'<p class="image-title">{title}</p>'
                    f'<img src="{image_data["url"]}" alt="{alt_text}">'
                    f'<p class="alt-text-display"><strong>Alt Text:</strong> {alt_text}</p>'
                    f'<p class="source-link"><a href="{image_data["source"]}" target="_blank">Source</a></p>'
                    f'</div>'
                )
                hybrid_content = hybrid_content.replace(original_placeholder, new_image_tag, 1)
            else:
                hybrid_content = hybrid_content.replace(original_placeholder, f'<p>[Image Placeholder: {title} - Could not fetch image]</p>', 1)
    else:
        # If images are disabled, just remove the placeholder
        hybrid_content = placeholder_regex.sub('', hybrid_content)


    final_html = markdown.markdown(hybrid_content, extensions=['fenced_code', 'tables'])
    return final_html

@retry_db_operation(max_retries=3)
def save_article_to_db(user_id, topic, content_html, content_raw, is_refined=False, article_id=None, chat_session_id=None):
    """Save article to database with retry logic"""
    try:
        # Extract word count (rough estimate)
        word_count = len(content_raw.split())

        if article_id:
            # Update existing article
            article = Article.query.filter_by(id=article_id, user_id=user_id).first()
            if article:
                article.content_html = content_html
                article.content_raw = content_raw
                article.is_refined = is_refined
                article.word_count = word_count
                article.updated_at = datetime.utcnow()
            else:
                # Article not found or doesn't belong to user
                return None
        else:
            # Create new article
            article = Article(
                user_id=user_id,
                title=topic[:200],  # Truncate if too long
                content_html=content_html,
                content_raw=content_raw,
                topic=topic,
                is_refined=is_refined,
                word_count=word_count,
                chat_session_id=chat_session_id  # Link to chat session
            )

            db.session.add(article)

            # Update user's article count and monthly word count
            user = User.query.get(user_id)
            if user:
                user.articles_generated = (user.articles_generated or 0) + 1

                # Check if we need to reset monthly quotas
                if user.last_quota_reset is None or (datetime.utcnow() - user.last_quota_reset) > timedelta(days=30):
                    user.words_generated_this_month = word_count
                    user.downloads_this_month = 0
                    user.last_quota_reset = datetime.utcnow()
                else:
                    user.words_generated_this_month = (user.words_generated_this_month or 0) + word_count

        db.session.commit()
        return article.id
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error saving article: {e}")
        return None

@retry_db_operation(max_retries=3)
def save_chat_session_to_db(user_id, session_id, title, messages, raw_text='', studio_type='ARTICLE'):
    """Save or update chat session to database"""
    try:
        # Check if chat session exists
        chat_session = ChatSession.query.filter_by(session_id=session_id, user_id=user_id).first()

        if chat_session:
            # Update existing session
            chat_session.title = title[:200]
            chat_session.set_messages(messages)
            chat_session.raw_text = raw_text
            chat_session.studio_type = studio_type
            chat_session.updated_at = datetime.utcnow()
        else:
            # Create new session
            chat_session = ChatSession(
                user_id=user_id,
                session_id=session_id,
                title=title[:200],
                raw_text=raw_text,
                studio_type=studio_type
            )
            chat_session.set_messages(messages)
            db.session.add(chat_session)

        db.session.commit()
        return chat_session.id
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error saving chat session: {e}")
        return None

@retry_db_operation(max_retries=2)
def check_monthly_word_quota(user):
    """Check if user has exceeded monthly word quota"""
    try:
        # Reset quota if it's been more than 30 days
        if user.last_quota_reset is None or (datetime.utcnow() - user.last_quota_reset) > timedelta(days=30):
            user.words_generated_this_month = 0
            user.downloads_this_month = 0
            user.last_quota_reset = datetime.utcnow()
            db.session.commit()
            return True

        words_generated = user.words_generated_this_month or 0
        return words_generated < MONTHLY_WORD_LIMIT
    except Exception as e:
        logger.error(f"Error checking word quota for user {user.id}: {e}")
        return True  # Allow operation if check fails

@retry_db_operation(max_retries=2)
def check_monthly_download_quota(user):
    """Check if user has exceeded monthly download quota"""
    try:
        # Reset quota if it's been more than 30 days
        if user.last_quota_reset is None or (datetime.utcnow() - user.last_quota_reset) > timedelta(days=30):
            user.words_generated_this_month = 0
            user.downloads_this_month = 0
            user.last_quota_reset = datetime.utcnow()
            db.session.commit()
            return True

        downloads_this_month = user.downloads_this_month or 0
        return downloads_this_month < MONTHLY_DOWNLOAD_LIMIT
    except Exception as e:
        logger.error(f"Error checking download quota for user {user.id}: {e}")
        return True  # Allow operation if check fails

def send_contact_email(name, email, subject, message):
    """Send contact form email"""
    try:
        if not app.config['MAIL_USERNAME'] or not app.config['MAIL_PASSWORD']:
            logger.warning("Email credentials not configured")
            return False

        msg = MIMEMultipart()
        msg['From'] = app.config['MAIL_USERNAME']
        msg['To'] = app.config['CONTACT_EMAIL']
        msg['Subject'] = f"InkDrive Contact Form: {subject}"

        body = f"""
New contact form submission from InkDrive:

Name: {name}
Email: {email}
Subject: {subject}

Message:
{message}

---
This message was sent from the InkDrive contact form.
        """

        msg.attach(MIMEText(body, 'plain'))

        server = smtplib.SMTP(app.config['MAIL_SERVER'], app.config['MAIL_PORT'])
        server.starttls()
        server.login(app.config['MAIL_USERNAME'], app.config['MAIL_PASSWORD'])
        text = msg.as_string()
        server.sendmail(app.config['MAIL_USERNAME'], app.config['CONTACT_EMAIL'], text)
        server.quit()

        return True
    except Exception as e:
        logger.error(f"Error sending contact email: {e}")
        return False

# --- 2. AUTHENTICATION ROUTES ---

@app.route('/auth/login', methods=['GET', 'POST'])
def auth_login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))

    form = LoginForm()
    if form.validate_on_submit():
        try:
            user = User.query.filter_by(email=form.email.data.lower()).first()

            if user:
                try:
                    password_valid = user.check_password(form.password.data)
                except Exception as e:
                    logger.error(f"Password check error for user {user.id}: {e}")
                    password_valid = False

                if password_valid:
                    login_user(user, remember=form.remember_me.data)

                    # Safely update last login with error handling
                    try:
                        user.update_last_login()
                    except Exception as e:
                        logger.warning(f"Failed to update last login for user {user.id}: {e}")

                    flash('Welcome back!', 'success')

                    next_page = request.args.get('next')
                    return redirect(next_page) if next_page else redirect(url_for('index'))
                else:
                    flash('Invalid email or password.', 'error')
            else:
                flash('Invalid email or password.', 'error')
        except (OperationalError, DatabaseError) as e:
            logger.error(f"Database error during login: {e}")
            flash('Database connection issue. Please try again in a moment.', 'error')
        except Exception as e:
            logger.error(f"Login error: {e}")
            flash('An error occurred during login. Please try again.', 'error')

    return render_template('auth/login.html', form=form)

@app.route('/auth/register', methods=['GET', 'POST'])
def auth_register():
    if current_user.is_authenticated:
        return redirect(url_for('index'))

    form = RegisterForm()
    if form.validate_on_submit():
        try:
            # Check if user already exists
            existing_user = User.query.filter_by(email=form.email.data.lower()).first()
            if existing_user:
                flash('Email address already registered.', 'error')
                return render_template('auth/register.html', form=form)

            # Create new user with safe defaults
            user = User(
                email=form.email.data.lower(),
                name=form.name.data,
                is_verified=True,  # Auto-verify for now
                is_active=True,
                words_generated_this_month=0,
                downloads_this_month=0,
                articles_generated=0,
                total_words_generated=0,
                last_quota_reset=datetime.utcnow()
            )

            # Set password with error handling
            try:
                user.set_password(form.password.data)
            except ValueError as e:
                flash('Invalid password format.', 'error')
                return render_template('auth/register.html', form=form)

            db.session.add(user)
            db.session.commit()

            login_user(user)
            flash('Registration successful! Welcome to InkDrive!', 'success')
            return redirect(url_for('index'))

        except (OperationalError, DatabaseError) as e:
            db.session.rollback()
            logger.error(f"Database error during registration: {e}")
            flash('Database connection issue. Please try again in a moment.', 'error')
        except Exception as e:
            db.session.rollback()
            logger.error(f"Registration error: {e}")
            flash('Registration failed. Please try again.', 'error')

    return render_template('auth/register.html', form=form)

@app.route('/auth/google')
def auth_google():
    """Initiate Google OAuth flow"""
    if not app.config.get('GOOGLE_CLIENT_ID'):
        flash('Google authentication is not configured.', 'error')
        return redirect(url_for('auth_login'))

    return render_template('auth/google_auth.html')

@app.route('/auth/google/callback', methods=['POST'])
def auth_google_callback():
    """Handle Google OAuth callback"""
    try:
        token = request.json.get('credential')
        if not token:
            return jsonify({'error': 'No credential provided'}), 400

        # Verify the token
        idinfo = id_token.verify_oauth2_token(
            token,
            google_requests.Request(),
            app.config['GOOGLE_CLIENT_ID']
        )

        if idinfo['iss'] not in ['accounts.google.com', 'https://accounts.google.com']:
            return jsonify({'error': 'Invalid token issuer'}), 400

        google_id = idinfo['sub']
        email = idinfo['email']
        name = idinfo['name']
        picture = idinfo.get('picture')

        # Check if user exists
        user = User.query.filter_by(google_id=google_id).first()
        if not user:
            user = User.query.filter_by(email=email.lower()).first()
            if user:
                # Link Google account to existing user
                user.google_id = google_id
                user.profile_picture = picture
            else:
                # Create new user with safe defaults
                user = User(
                    email=email.lower(),
                    name=name,
                    google_id=google_id,
                    profile_picture=picture,
                    is_verified=True,
                    is_active=True,
                    words_generated_this_month=0,
                    downloads_this_month=0,
                    articles_generated=0,
                    total_words_generated=0,
                    last_quota_reset=datetime.utcnow()
                )
                db.session.add(user)
        else:
            # Update existing Google user info
            user.name = name
            user.profile_picture = picture

        db.session.commit()
        login_user(user)

        # Safely update last login
        try:
            user.update_last_login()
        except Exception as e:
            logger.warning(f"Failed to update last login for Google user {user.id}: {e}")

        return jsonify({'success': True, 'redirect': url_for('index')})

    except ValueError as e:
        logger.error(f"Google auth token error: {e}")
        return jsonify({'error': 'Invalid token'}), 400
    except (OperationalError, DatabaseError) as e:
        db.session.rollback()
        logger.error(f"Database error during Google auth: {e}")
        return jsonify({'error': 'Database connection issue. Please try again.'}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Google auth error: {e}")
        return jsonify({'error': 'Authentication failed'}), 500

@app.route('/auth/logout')
@login_required
def auth_logout():
    logout_user()
    flash('You have been logged out.', 'info')
    return redirect(url_for('index'))

# --- 3. PROFILE & DASHBOARD ROUTES ---

@app.route('/profile')
@app.route('/profile/dashboard')
@login_required
def profile_dashboard():
    """Redirects to the main studios page, as this dashboard is deprecated."""
    return redirect(url_for('index'))

@app.route('/profile/edit', methods=['GET', 'POST'])
@login_required
def profile_edit():
    """Edit user profile"""
    form = ProfileForm(obj=current_user)

    if form.validate_on_submit():
        try:
            current_user.name = form.name.data
            current_user.theme_preference = form.theme_preference.data

            db.session.commit()
            flash('Profile updated successfully!', 'success')
            return redirect(url_for('profile_dashboard'))
        except (OperationalError, DatabaseError) as e:
            db.session.rollback()
            logger.error(f"Database error updating profile: {e}")
            flash('Database connection issue. Please try again.', 'error')
        except Exception as e:
            db.session.rollback()
            logger.error(f"Profile update error: {e}")
            flash('Failed to update profile.', 'error')

    return render_template('profile/edit.html', form=form)

@app.route('/profile/change-password', methods=['GET', 'POST'])
@login_required
def profile_change_password():
    """Change user password"""
    if current_user.google_id and not current_user.password_hash:
        flash('Google account users cannot change password here.', 'info')
        return redirect(url_for('profile_dashboard'))

    form = ChangePasswordForm()

    if form.validate_on_submit():
        try:
            if not current_user.check_password(form.current_password.data):
                flash('Current password is incorrect.', 'error')
                return render_template('profile/change_password.html', form=form)

            current_user.set_password(form.new_password.data)
            db.session.commit()
            flash('Password changed successfully!', 'success')
            return redirect(url_for('profile_dashboard'))
        except (OperationalError, DatabaseError) as e:
            db.session.rollback()
            logger.error(f"Database error changing password: {e}")
            flash('Database connection issue. Please try again.', 'error')
        except Exception as e:
            db.session.rollback()
            logger.error(f"Password change error: {e}")
            flash('Failed to change password.', 'error')

    return render_template('profile/change_password.html', form=form)

@app.route('/profile/delete-account', methods=['POST'])
@login_required
def profile_delete_account():
    """Delete user account"""
    try:
        # Delete all user's articles
        Article.query.filter_by(user_id=current_user.id).delete()

        # Delete all user's chat sessions
        ChatSession.query.filter_by(user_id=current_user.id).delete()

        # Delete the user
        db.session.delete(current_user)
        db.session.commit()

        logout_user()
        flash('Your account has been deleted successfully.', 'success')
        return redirect(url_for('index'))
    except (OperationalError, DatabaseError) as e:
        db.session.rollback()
        logger.error(f"Database error deleting account: {e}")
        flash('Database connection issue. Please try again.', 'error')
        return redirect(url_for('profile_edit'))
    except Exception as e:
        db.session.rollback()
        logger.error(f"Account deletion error: {e}")
        flash('Failed to delete account. Please try again.', 'error')
        return redirect(url_for('profile_edit'))



@app.route('/article/<int:article_id>/delete', methods=['POST'])
@login_required
def delete_article(article_id):
    """Delete an article and its associated chat session"""
    try:
        article = Article.query.filter_by(id=article_id, user_id=current_user.id).first_or_404()

        # Find and delete associated chat session if it exists
        if article.chat_session_id:
            chat_session = ChatSession.query.filter_by(id=article.chat_session_id, user_id=current_user.id).first()
            if chat_session:
                db.session.delete(chat_session)

        # Delete the article
        db.session.delete(article)
        db.session.commit()

        return jsonify({
            "success": True,
            "message": "Article and associated chat deleted successfully!"
        })
    except (OperationalError, DatabaseError) as e:
        db.session.rollback()
        logger.error(f"Database error deleting article {article_id}: {e}")
        return jsonify({"error": "Database connection issue"}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Article delete error: {e}")
        return jsonify({"error": "Failed to delete article"}), 500


# --- 4. MAIN APP ROUTES ---

@app.route("/")
def index():
    """Main application route"""
    if current_user.is_authenticated:
        return render_template("dashboard.html", user=current_user, page_type='dashboard')
    else:
        return render_template("landing.html", featured_articles=[])

@app.route('/studio/article')
@login_required
def article_studio():
    """The new Article Studio page"""
    return render_template('article_studio.html', user=current_user, page_type='studio', studio_type='article')

@app.route('/studio/social')
@login_required
def social_studio():
    """The new Social & Comms Studio page"""
    return render_template('social_studio.html', user=current_user, page_type='studio', studio_type='social')

@app.route('/studio/editing')
@login_required
def editing_studio():
    """The new Editing & Refinement Studio page"""
    return render_template('editing_studio.html', user=current_user, page_type='studio', studio_type='editing')

@app.route('/studio/repurpose')
@login_required
def repurpose_studio():
    """The new Content Repurposing Studio page"""
    return render_template('repurposing_studio.html', user=current_user, page_type='studio', studio_type='repurpose')

@app.route('/studio/seo')
@login_required
def seo_studio():
    """The new SEO Strategy Studio page"""
    return render_template('seo_studio.html', user=current_user, page_type='studio', studio_type='seo')

@app.route('/studio/brainstorming')
@login_required
def brainstorming_studio():
    """The new Brainstorming Studio page"""
    return render_template('brainstorming_studio.html', user=current_user, page_type='studio', studio_type='brainstorming')

@app.route('/studio/scriptwriting')
@login_required
def scriptwriting_studio():
    """The new Scriptwriting Studio page"""
    return render_template('scriptwriting_studio.html', user=current_user, page_type='studio', studio_type='scriptwriting')

@app.route('/studio/ecommerce')
@login_required
def ecommerce_studio():
    """The new E-commerce Studio page"""
    return render_template('ecommerce_studio.html', user=current_user, page_type='studio', studio_type='ecommerce')

@app.route('/studio/webcopy')
@login_required
def webcopy_studio():
    """The new Web Copy Studio page"""
    return render_template('webcopy_studio.html', user=current_user, page_type='studio', studio_type='webcopy')

@app.route('/studio/business')
@login_required
def business_studio():
    """The new Business Docs Studio page"""
    return render_template('business_studio.html', user=current_user, page_type='studio', studio_type='business')

@app.route('/api/v1/generate/social', methods=['POST'])
@login_required
def generate_social_content():
    """Handles various social and communication content generation requests."""
    if not CLIENT:
        return jsonify({"error": "AI service is not available."}), 503

    data = request.get_json()
    tool = data.get("tool")
    settings = data.get("settings", {})
    chat_session_id = data.get("chat_session_id")

    if not tool:
        return jsonify({"error": "Missing required field: tool."}), 400

    if not check_monthly_word_quota(current_user):
        return jsonify({"error": f"You've reached your monthly word limit."}), 403

    if tool == 'social_post':
        topic = data.get("topic")
        goal = data.get("goal")
        platform = data.get("platform")
        if not all([topic, goal, platform]):
            return jsonify({"error": "Missing required fields for social post."}), 400

        full_prompt = construct_social_post_prompt(topic, goal, platform, settings)
        try:
            response = CLIENT.generate_content(contents=full_prompt)
            raw_text = response.candidates[0].content.parts[0].text
            posts = [p.strip() for p in raw_text.split('---') if p.strip()]

            if not chat_session_id:
                chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"
            user_message = json.dumps({"tool": "social_post", "topic": topic, "goal": goal, "platform": platform, "settings": settings})
            messages = [{"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"}, {"content": raw_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}]
            session_title = f"Social Posts for '{topic}'"
            save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, raw_text, studio_type='SOCIAL_POST')

            return jsonify({"posts": posts, "chat_session_id": chat_session_id})
        except Exception as e:
            logger.error(f"Social post generation error: {e}")
            return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

    elif tool == 'email':
        topic = data.get("topic")
        audience = data.get("audience")
        tone = data.get("tone")
        if not all([topic, audience, tone]):
            return jsonify({"error": "Missing required fields for email."}), 400

        full_prompt = construct_email_prompt(topic, audience, tone, settings)
        try:
            response = CLIENT.generate_content(contents=full_prompt)
            raw_text = response.candidates[0].content.parts[0].text

            if not chat_session_id:
                chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"
            user_message = json.dumps({"tool": "email", "topic": topic, "audience": audience, "tone": tone, "settings": settings})
            messages = [{"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"}, {"content": raw_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}]
            session_title = f"Email for '{topic}'"
            save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, raw_text, studio_type='EMAIL')

            return jsonify({"email_content": raw_text, "chat_session_id": chat_session_id})
        except Exception as e:
            logger.error(f"Email generation error: {e}")
            return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

    elif tool == 'ad_copy':
        product = data.get("product")
        audience = data.get("audience")
        if not all([product, audience]):
            return jsonify({"error": "Missing required fields for ad copy."}), 400

        full_prompt = construct_ad_copy_prompt(product, audience, settings)
        try:
            response = CLIENT.generate_content(contents=full_prompt)
            raw_text = response.candidates[0].content.parts[0].text
            ad_copy = [ad.strip() for ad in raw_text.split('---') if ad.strip()]

            if not chat_session_id:
                chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"
            user_message = json.dumps({"tool": "ad_copy", "product": product, "audience": audience, "settings": settings})
            messages = [{"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"}, {"content": raw_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}]
            session_title = f"Ad Copy for {product}"
            save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, raw_text, studio_type='AD_COPY')

            return jsonify({"ad_copy": ad_copy, "chat_session_id": chat_session_id})
        except Exception as e:
            logger.error(f"Ad copy generation error: {e}")
            return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

    else:
        return jsonify({"error": f"Unknown tool: {tool}"}), 400

@app.route('/api/v1/generate/brainstorm', methods=['POST'])
@login_required
def generate_brainstorm_content():
    """Handles various brainstorming and naming requests."""
    if not CLIENT:
        return jsonify({"error": "AI service is not available."}), 503

    data = request.get_json()
    tool = data.get("tool")
    settings = data.get("settings", {})
    chat_session_id = data.get("chat_session_id")

    if not tool:
        return jsonify({"error": "Missing required field: tool."}), 400

    if not check_monthly_word_quota(current_user):
        return jsonify({"error": f"You've reached your monthly word limit."}), 403

    if tool == 'general':
        if not settings.get('goal'):
            return jsonify({"error": "Missing required field: goal."}), 400

        full_prompt = construct_brainstorm_prompt(settings)
        try:
            response = CLIENT.generate_content(contents=full_prompt)
            raw_text = response.candidates[0].content.parts[0].text
            ideas = [idea.strip() for idea in raw_text.split('\n') if idea.strip()]

            if not chat_session_id:
                chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"
            user_message = json.dumps({"tool": "general", "settings": settings})
            messages = [{"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"}, {"content": raw_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}]
            session_title = f"Brainstorm: {settings.get('goal', 'General')}"
            save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, raw_text, studio_type='IDEAS')

            return jsonify({"ideas": ideas, "chat_session_id": chat_session_id})
        except Exception as e:
            logger.error(f"Brainstorming error: {e}")
            return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

    elif tool == 'naming':
        if not settings.get('description'):
            return jsonify({"error": "Missing required field: description."}), 400

        full_prompt = construct_naming_prompt(settings)
        try:
            response = CLIENT.generate_content(contents=full_prompt)
            raw_text = response.candidates[0].content.parts[0].text
            names = [name.strip() for name in raw_text.split('\n') if name.strip()]

            if not chat_session_id:
                chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"
            user_message = json.dumps({"tool": "naming", "settings": settings})
            messages = [{"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"}, {"content": raw_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}]
            session_title = f"Names for: {settings.get('description', 'New Project')[:30]}"
            save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, raw_text, studio_type='NAMING')

            return jsonify({"names": names, "chat_session_id": chat_session_id})
        except Exception as e:
            logger.error(f"Naming error: {e}")
            return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

    else:
        return jsonify({"error": f"Unknown tool: {tool}"}), 400

@app.route('/api/v1/generate/script', methods=['POST'])
@login_required
def generate_script():
    """Generates a script based on detailed user settings."""
    if not CLIENT:
        return jsonify({"error": "AI service is not available."}), 503

    data = request.get_json()
    topic = data.get("topic")
    settings = data.get("settings", {})
    chat_session_id = data.get("chat_session_id")

    if not topic:
        return jsonify({"error": "Missing required field: topic."}), 400

    if not check_monthly_word_quota(current_user):
        return jsonify({"error": f"You've reached your monthly word limit."}), 403

    full_prompt = construct_script_prompt(topic, settings)
    try:
        response = CLIENT.generate_content(contents=full_prompt)
        script_text = response.candidates[0].content.parts[0].text

        # Save to chat history
        if not chat_session_id:
            chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"

        # Store the original user request (topic and all settings) for session restoration
        user_message = json.dumps({"topic": topic, "settings": settings})
        messages = [
            {"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"},
            {"content": script_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}
        ]

        session_title = f"Script: {topic[:40]}{'...' if len(topic) > 40 else ''}"
        save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, script_text, studio_type='SCRIPT')

        return jsonify({
            "script": script_text,
            "chat_session_id": chat_session_id
        })
    except Exception as e:
        logger.error(f"Script generation error: {e}")
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

@app.route('/api/v1/generate/ecommerce', methods=['POST'])
@login_required
def generate_ecommerce():
    """Handles various e-commerce content generation requests."""
    if not CLIENT:
        return jsonify({"error": "AI service is not available."}), 503

    data = request.get_json()
    tool = data.get("tool")
    settings = data.get("settings", {})
    chat_session_id = data.get("chat_session_id")

    if not tool:
        return jsonify({"error": "Missing required field: tool."}), 400

    if not check_monthly_word_quota(current_user):
        return jsonify({"error": f"You've reached your monthly word limit."}), 403

    # Route to the correct prompt constructor based on the tool
    if tool == 'description':
        full_prompt = construct_product_description_prompt(settings)
        session_title = f"Desc: {settings.get('productName', 'New Product')[:30]}"
    elif tool == 'campaign':
        full_prompt = construct_campaign_prompt(settings)
        session_title = f"Campaign: {settings.get('occasion', 'New Event')[:30]}"
    elif tool == 'review':
        full_prompt = construct_review_response_prompt(settings)
        session_title = f"Review Resp: {settings.get('review', 'New Review')[:30]}"
    else:
        return jsonify({"error": f"Unknown tool: {tool}"}), 400

    try:
        response = CLIENT.generate_content(contents=full_prompt)
        result_text = response.candidates[0].content.parts[0].text

        # Save to chat history
        if not chat_session_id:
            chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"

        user_message = json.dumps({"tool": tool, "settings": settings})
        messages = [
            {"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"},
            {"content": result_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}
        ]

        save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, result_text, studio_type='ECOMMERCE')

        return jsonify({
            "result": result_text,
            "chat_session_id": chat_session_id
        })
    except Exception as e:
        logger.error(f"E-commerce generation error for tool {tool}: {e}")
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

@app.route('/api/v1/generate/webcopy', methods=['POST'])
@login_required
def generate_webcopy():
    """Handles various web copy generation requests."""
    if not CLIENT:
        return jsonify({"error": "AI service is not available."}), 503

    data = request.get_json()
    tool = data.get("tool")
    settings = data.get("settings", {})
    chat_session_id = data.get("chat_session_id")

    if not tool:
        return jsonify({"error": "Missing required field: tool."}), 400

    if current_user.is_authenticated:
        if not check_monthly_word_quota(current_user):
            return jsonify({"error": f"You've reached your monthly word limit."}), 403

    # Route to the correct prompt constructor
    if tool == 'landing_page':
        full_prompt = construct_landing_page_prompt(settings)
        session_title = f"LP: {settings.get('productName', 'New Page')[:30]}"
    elif tool == 'homepage_section':
        full_prompt = construct_homepage_section_prompt(settings)
        session_title = f"Homepage: {settings.get('sectionType', 'New Section')[:30]}"
    elif tool == 'usp':
        full_prompt = construct_usp_prompt(settings)
        session_title = f"USP for: {settings.get('productDesc', 'New Product')[:30]}"
    else:
        return jsonify({"error": f"Unknown tool: {tool}"}), 400

    try:
        response = CLIENT.generate_content(contents=full_prompt)
        result_text = response.candidates[0].content.parts[0].text

        if not chat_session_id:
            chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"

        user_message = json.dumps({"tool": tool, "settings": settings})
        messages = [
            {"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"},
            {"content": result_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}
        ]

        save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, result_text, studio_type='WEBCOPY')

        return jsonify({
            "result": result_text,
            "chat_session_id": chat_session_id
        })
    except Exception as e:
        logger.error(f"Web copy generation error for tool {tool}: {e}")
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

@app.route('/api/v1/generate/business', methods=['POST'])
# @login_required
def generate_business_doc():
    """Handles various business document generation requests."""
    if not CLIENT:
        return jsonify({"error": "AI service is not available."}), 503

    data = request.get_json()
    tool = data.get("tool")
    settings = data.get("settings", {})
    chat_session_id = data.get("chat_session_id")

    if not tool:
        return jsonify({"error": "Missing required field: tool."}), 400

    if not check_monthly_word_quota(current_user):
        return jsonify({"error": f"You've reached your monthly word limit."}), 403

    # Route to the correct prompt constructor
    if tool == 'proposal':
        full_prompt = construct_proposal_prompt(settings)
        session_title = f"Proposal for {settings.get('client', 'New Client')}"
    elif tool == 'report':
        full_prompt = construct_report_prompt(settings)
        session_title = f"Report: {settings.get('subject', 'New Report')[:30]}"
    elif tool == 'press_release':
        full_prompt = construct_press_release_prompt(settings)
        session_title = f"PR: {settings.get('headline', 'New Release')[:30]}"
    else:
        return jsonify({"error": f"Unknown tool: {tool}"}), 400

    try:
        response = CLIENT.generate_content(contents=full_prompt)
        result_text = response.candidates[0].content.parts[0].text

        if not chat_session_id:
            chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"

        user_message = json.dumps({"tool": tool, "settings": settings})
        messages = [
            {"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"},
            {"content": result_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}
        ]

        # Note: The studio_type is generic 'BUSINESS' now, but could be made more specific
        if current_user.is_authenticated:
            save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, result_text, studio_type='BUSINESS')

        return jsonify({
            "result": result_text,
            "chat_session_id": chat_session_id
        })
    except Exception as e:
        logger.error(f"Business doc generation error for tool {tool}: {e}")
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

@app.route("/api/v1/generate/article", methods=["POST"])
@login_required
def generate_article():
    if not CLIENT:
        return jsonify({"error": "AI service is not available."}), 503

    data = request.get_json()
    user_topic = data.get("topic")
    settings = data.get("settings", {})
    chat_session_id = data.get("chat_session_id")

    logger.info(f"Generating article for topic: {user_topic}")
    logger.info(f"Settings: {settings}")
    if not user_topic:
        return jsonify({"error": "Topic is missing."}), 400

    # Check monthly word quota
    if not check_monthly_word_quota(current_user):
        logger.warning(f"User {current_user.id} has reached their word limit.")
        return jsonify({"error": f"You've reached your monthly limit of {MONTHLY_WORD_LIMIT} words. Please try again next month."}), 403

    enable_images = settings.get('enable_images', False) and current_user.is_superadmin
    logger.info(f"Image generation enabled: {enable_images}")

    full_prompt = construct_initial_prompt(user_topic, settings)
    try:
        response = CLIENT.generate_content(contents=full_prompt)
        logger.info("Successfully generated content from AI client.")

        if not response.candidates or not response.candidates[0].content.parts:
            logger.error("Model response was empty or blocked.")
            return jsonify({"error": "Model response was empty or blocked."}), 500

        raw_text = response.candidates[0].content.parts[0].text
        final_html = format_article_content(raw_text, user_topic, enable_images)
        logger.info("Successfully formatted article content.")

        # Save chat session to database
        if not chat_session_id:
            chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"

        messages = [
            {"content": user_topic, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"},
            {"content": final_html, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}
        ]

        db_chat_session_id = save_chat_session_to_db(current_user.id, chat_session_id, user_topic, messages, raw_text, studio_type='ARTICLE')

        # Save article to database with chat session link
        article_id = save_article_to_db(current_user.id, user_topic, final_html, raw_text, False, None, db_chat_session_id)

        return jsonify({
            "article_html": final_html,
            "raw_text": raw_text,
            "article_id": article_id,
            "chat_session_id": chat_session_id,
            "refinements_remaining": 5  # Always 5 for authenticated users on new article
        })
    except Exception as e:
        logger.error(f"Article generation error: {e}")
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

@app.route("/api/v1/generate/article-guest", methods=["POST"])
def generate_guest_article():
    """Generate article for guest users"""
    if not CLIENT:
        return jsonify({"error": "AI service is not available."}), 503

    data = request.get_json()
    user_topic = data.get("topic")
    if not user_topic:
        return jsonify({"error": "Topic is missing."}), 400

    full_prompt = construct_initial_prompt(user_topic)
    try:
        response = CLIENT.generate_content(contents=full_prompt)

        if not response.candidates or not response.candidates[0].content.parts:
            return jsonify({"error": "Model response was empty or blocked."}), 500

        raw_text = response.candidates[0].content.parts[0].text
        final_html = format_article_content(raw_text, user_topic)

        # For guests, we don't save to database
        return jsonify({
            "article_html": final_html,
            "raw_text": raw_text,
            "refinements_remaining": 1  # Only 1 for guests
        })
    except Exception as e:
        logger.error(f"Guest article generation error: {e}")
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

@app.route("/api/v1/refine/article", methods=["POST"])
def refine_article():
    """Refine article for both authenticated and guest users"""
    if not CLIENT:
        return jsonify({"error": "AI service is not available."}), 503

    data = request.get_json()
    raw_text, refinement_prompt = data.get("raw_text"), data.get("refinement_prompt")
    article_id = data.get("article_id")
    chat_session_id = data.get("chat_session_id")
    refinements_used = data.get("refinements_used", 0)
    topic = data.get("topic", "")

    if not all([raw_text, refinement_prompt]):
        return jsonify({"error": "Missing data for refinement."}), 400

    # Check refinement limits and word quota for authenticated users
    if current_user.is_authenticated:
        if refinements_used >= 5:
            return jsonify({"error": "You've used all 5 refinements for this article."}), 403

        if not check_monthly_word_quota(current_user):
            return jsonify({"error": f"You've reached your monthly limit of {MONTHLY_WORD_LIMIT} words. Please try again next month."}), 403
    else:
        # Guest users get only 1 refinement
        if refinements_used >= 1:
            return jsonify({"error": "You've used your 1 refinement. Sign up for a free account to get 5 refinements per article."}), 403

    try:
        history = [
            {"role": "user", "parts": [{"text": "You are an AI assistant. You provided this article draft."}]},
            {"role": "model", "parts": [{"text": raw_text}]},
            {"role": "user", "parts": [{"text": refinement_prompt}]}
        ]
        response = CLIENT.generate_content(contents=history)

        if not response.candidates or not response.candidates[0].content.parts:
            return jsonify({"error": "Refinement response was empty or blocked."}), 500

        refined_text = response.candidates[0].content.parts[0].text
        final_html = format_article_content(refined_text, topic)

        # Update word count for authenticated users
        if current_user.is_authenticated:
            word_count = len(refined_text.split())
            current_user.words_generated_this_month = (current_user.words_generated_this_month or 0) + word_count
            db.session.commit()

            # Update article in database if article_id provided
            if article_id:
                save_article_to_db(current_user.id, "", final_html, refined_text, True, article_id)

            # Update chat session
            if chat_session_id:
                chat_session = ChatSession.query.filter_by(session_id=chat_session_id, user_id=current_user.id).first()
                if chat_session:
                    new_user_message = {"content": refinement_prompt, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"}
                    new_ai_message = {"content": final_html, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}

                    # Get current messages as a string
                    current_messages_str = chat_session.messages or '[]'

                    # Create string for new messages
                    new_user_message_str = json.dumps(new_user_message)
                    new_ai_message_str = json.dumps(new_ai_message)

                    # Append new messages to the string
                    if current_messages_str == '[]':
                        updated_messages_str = f"[{new_user_message_str},{new_ai_message_str}]"
                    else:
                        # Remove the closing ']' from the current messages
                        current_messages_trimmed = current_messages_str[:-1]
                        updated_messages_str = f"{current_messages_trimmed},{new_user_message_str},{new_ai_message_str}]"

                    chat_session.messages = updated_messages_str
                    chat_session.raw_text = refined_text
                    chat_session.has_refined = True
                    db.session.commit()

        new_refinements_used = refinements_used + 1
        remaining_refinements = (5 if current_user.is_authenticated else 1) - new_refinements_used

        return jsonify({
            "article_html": final_html,
            "raw_text": refined_text,
            "refinements_used": new_refinements_used,
            "refinements_remaining": remaining_refinements
        })
    except Exception as e:
        logger.error(f"Article refinement error: {e}")
        return jsonify({"error": f"An unexpected error occurred during refinement: {str(e)}"}), 500

@app.route('/api/v1/refine/text', methods=['POST'])
@login_required
def refine_and_edit_text():
    """Handles various text refinement and editing requests."""
    if not CLIENT:
        return jsonify({"error": "AI service is not available."}), 503

    data = request.get_json()
    tool = data.get("tool")
    settings = data.get("settings", {})
    chat_session_id = data.get("chat_session_id")
    text = data.get("text")

    if not tool or not text:
        return jsonify({"error": "Missing required fields: tool, text."}), 400

    if not check_monthly_word_quota(current_user):
        return jsonify({"error": f"You've reached your monthly word limit."}), 403

    if tool == 'tone_style':
        full_prompt = construct_refine_text_prompt(text, settings)
        try:
            response = CLIENT.generate_content(contents=full_prompt)
            refined_text = response.candidates[0].content.parts[0].text

            if not chat_session_id:
                chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"
            user_message = json.dumps({"tool": "tone_style", "text": text, "settings": settings})
            messages = [{"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"}, {"content": refined_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}]
            session_title = f"Refinement: {settings.get('goal', 'General')}"
            save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, refined_text, studio_type='TEXT_REFINEMENT')

            return jsonify({"refined_text": refined_text, "chat_session_id": chat_session_id})
        except Exception as e:
            logger.error(f"Text refinement error: {e}")
            return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

    elif tool == 'summarize':
        full_prompt = construct_summarizer_prompt(text, settings)
        try:
            response = CLIENT.generate_content(contents=full_prompt)
            summary = response.candidates[0].content.parts[0].text

            if not chat_session_id:
                chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"
            user_message = json.dumps({"tool": "summarize", "text": text, "settings": settings})
            messages = [{"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"}, {"content": summary, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}]
            session_title = f"Summary ({settings.get('format', 'paragraph')})"
            save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, summary, studio_type='SUMMARY')

            return jsonify({"summary": summary, "chat_session_id": chat_session_id})
        except Exception as e:
            logger.error(f"Summarization error: {e}")
            return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

    elif tool == 'translate':
        full_prompt = construct_translator_prompt(text, settings)
        try:
            response = CLIENT.generate_content(contents=full_prompt)
            translated_text = response.candidates[0].content.parts[0].text

            if not chat_session_id:
                chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"
            user_message = json.dumps({"tool": "translate", "text": text, "settings": settings})
            messages = [{"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"}, {"content": translated_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}]
            session_title = f"Translation to {settings.get('locale', 'Unknown')}"
            save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, translated_text, studio_type='TRANSLATION')

            return jsonify({"translated_text": translated_text, "chat_session_id": chat_session_id})
        except Exception as e:
            logger.error(f"Translation error: {e}")
            return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

    else:
        return jsonify({"error": f"Unknown tool: {tool}"}), 400

@app.route('/api/v1/repurpose/content', methods=['POST'])
@login_required
def repurpose_content():
    """Handles various content repurposing requests."""
    if not CLIENT:
        return jsonify({"error": "AI service is not available."}), 503

    data = request.get_json()
    tool = data.get("tool")
    text = data.get("text")
    settings = data.get("settings", {})
    chat_session_id = data.get("chat_session_id")

    if not all([tool, text]):
        return jsonify({"error": "Missing required fields: tool or text."}), 400

    if not check_monthly_word_quota(current_user):
        return jsonify({"error": f"You've reached your monthly word limit."}), 403

    full_prompt = construct_repurpose_prompt(tool, text, settings)
    try:
        response = CLIENT.generate_content(contents=full_prompt)
        result_text = response.candidates[0].content.parts[0].text

        if not chat_session_id:
            chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"

        # Determine the studio type for saving the session
        studio_type_map = {
            "twitter_thread": "REPURPOSE_TWEET",
            "linkedin_post": "REPURPOSE_LINKEDIN", # Assuming a new type
            "video_script": "REPURPOSE_SLIDES" # As per user clarification
        }
        studio_type = studio_type_map.get(tool, "REPURPOSE")

        user_message = json.dumps({"tool": tool, "text": text, "settings": settings})
        messages = [{"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"}, {"content": result_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}]
        session_title = f"Repurposed into: {tool.replace('_', ' ').title()}"
        save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, result_text, studio_type=studio_type)

        return jsonify({"result": result_text, "chat_session_id": chat_session_id})
    except Exception as e:
        logger.error(f"Content repurposing error for tool {tool}: {e}")
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

@app.route('/api/v1/seo/tools', methods=['POST'])
@login_required
def seo_tools():
    """Handles various SEO tool requests."""
    if not CLIENT:
        return jsonify({"error": "AI service is not available."}), 503

    data = request.get_json()
    tool = data.get("tool")
    settings = data.get("settings", {})
    chat_session_id = data.get("chat_session_id")

    if not tool:
        return jsonify({"error": "Missing required field: tool."}), 400

    if not check_monthly_word_quota(current_user):
        return jsonify({"error": f"You've reached your monthly word limit."}), 403

    if tool == 'keyword_strategy':
        if not settings.get('description') or not settings.get('audience'):
            return jsonify({"error": "Missing required fields for keyword strategy."}), 400

        full_prompt = construct_keyword_strategy_prompt(settings)
        try:
            response = CLIENT.generate_content(contents=full_prompt)
            raw_text = response.candidates[0].content.parts[0].text
            keywords_json = json.loads(raw_text)

            if not chat_session_id:
                chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"
            user_message = json.dumps({"tool": "keyword_strategy", "settings": settings})
            messages = [{"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"}, {"content": raw_text, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}]
            session_title = "Keyword Strategy"
            save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, raw_text, studio_type='SEO_KEYWORDS')

            return jsonify({"keywords": keywords_json, "chat_session_id": chat_session_id})
        except json.JSONDecodeError:
            logger.error(f"SEO keywords JSON parsing error. Raw text: {raw_text}")
            return jsonify({"error": "Failed to parse the keyword data from the AI. Please try again."}), 500
        except Exception as e:
            logger.error(f"Keyword strategy error: {e}")
            return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

    elif tool == 'on_page_audit':
        if not settings.get('url') or not settings.get('primaryKeyword'):
            return jsonify({"error": "Missing required fields for on-page audit."}), 400

        full_prompt = construct_seo_audit_prompt(settings)
        try:
            response = CLIENT.generate_content(contents=full_prompt)
            audit_results = response.candidates[0].content.parts[0].text

            if not chat_session_id:
                chat_session_id = f"chat_{int(datetime.utcnow().timestamp())}_{current_user.id}"
            user_message = json.dumps({"tool": "on_page_audit", "settings": settings})
            messages = [{"content": user_message, "isUser": True, "id": f"msg_{int(datetime.utcnow().timestamp())}_user"}, {"content": audit_results, "isUser": False, "id": f"msg_{int(datetime.utcnow().timestamp())}_ai"}]
            session_title = f"On-Page Audit for {settings.get('primaryKeyword')}"
            save_chat_session_to_db(current_user.id, chat_session_id, session_title, messages, audit_results, studio_type='SEO_AUDIT')

            return jsonify({"audit_results": audit_results, "chat_session_id": chat_session_id})
        except Exception as e:
            logger.error(f"On-page audit error: {e}")
            return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500

    else:
        return jsonify({"error": f"Unknown tool: {tool}"}), 400

@app.route("/download-docx", methods=["POST"])
def download_docx():
    """Download article as DOCX for both authenticated and guest users"""
    data = request.get_json()
    html_content = data.get("html")
    topic = data.get("topic", "Generated Article")
    article_id = data.get("article_id")

    if not html_content:
        return jsonify({"error": "Missing HTML content."}), 400

    try:
        # Update download count if article_id provided and user is authenticated
        if article_id and current_user.is_authenticated:
            try:
                # Check monthly download quota
                if not check_monthly_download_quota(current_user):
                    return jsonify({"error": f"You've reached your monthly limit of {MONTHLY_DOWNLOAD_LIMIT} downloads."}), 403

                article = Article.query.filter_by(id=article_id, user_id=current_user.id).first()
                if article:
                    article.increment_download()
                    current_user.downloads_this_month = (getattr(current_user, 'downloads_this_month', 0) or 0) + 1
                    db.session.commit()
            except Exception as e:
                logger.warning(f"Failed to update download count: {e}")

        # Generate DOCX
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()
        doc.add_heading(topic, level=0)

        for element in soup.find_all(['h2', 'h3', 'p', 'div']):
            if element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'p' and not element.find_parents("div"):
                doc.add_paragraph(element.get_text())
            elif element.name == 'div' and "real-image-container" in element.get('class', []):
                title_p = element.find('p', class_='image-title')
                img_tag = element.find('img')
                alt_text_p = element.find('p', class_='alt-text-display')
                attr_p = element.find('p', class_='attribution')

                if title_p:
                    p = doc.add_paragraph(title_p.get_text())
                    p.alignment = 1
                    p.bold = True

                if img_tag and img_tag.get('src'):
                    try:
                        img_response = requests.get(img_tag['src'], stream=True, timeout=10)
                        img_response.raise_for_status()
                        doc.add_picture(io.BytesIO(img_response.content), width=Inches(5.5))
                    except requests.RequestException:
                        doc.add_paragraph(f"[Image failed to load from {img_tag['src']}]")

                if alt_text_p:
                    p = doc.add_paragraph()
                    clean_alt_text = alt_text_p.get_text()
                    if clean_alt_text.lower().startswith("alt text:"):
                        clean_alt_text = clean_alt_text[len("Alt Text:"):].strip()
                    run = p.add_run(clean_alt_text)
                    run.italic = True
                    p.alignment = 1

                if attr_p:
                    p = doc.add_paragraph(attr_p.get_text())
                    p.alignment = 1
                    p.italic = True

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        filename = f"{topic[:50].strip().replace(' ', '_')}.docx"
        return send_file(file_stream, as_attachment=True, download_name=filename,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        logger.error(f"DOCX generation error: {e}")
        return jsonify({"error": "Failed to generate document."}), 500

# --- 5. API ROUTES FOR USER DATA ---

@app.route('/api/user/chat-history')
@login_required
def api_user_chat_history():
    """Get user's chat history from database, with optional studio filter."""
    try:
        studio_filter = request.args.get('studio') # e.g., 'social', 'seo'

        query = ChatSession.query.filter_by(user_id=current_user.id)

        if studio_filter and studio_filter in STUDIO_TYPE_MAPPING:
            types_to_filter = STUDIO_TYPE_MAPPING[studio_filter]
            query = query.filter(ChatSession.studio_type.in_(types_to_filter))

        chat_sessions = query.order_by(ChatSession.updated_at.desc()).all()

        return jsonify([session.to_dict() for session in chat_sessions])
    except (OperationalError, DatabaseError) as e:
        logger.error(f"Database error in API chat history: {e}")
        return jsonify({"error": "Database connection issue"}), 500
    except Exception as e:
        logger.error(f"API chat history error: {e}")
        return jsonify({"error": "Failed to fetch chat history"}), 500

@app.route('/api/chat-session/<string:session_id>')
@login_required
def api_get_chat_session(session_id):
    """Get a specific chat session"""
    try:
        chat_session = ChatSession.query.filter_by(session_id=session_id, user_id=current_user.id).first()
        if not chat_session:
            return jsonify({"error": "Chat session not found"}), 404

        return jsonify(chat_session.to_dict())
    except (OperationalError, DatabaseError) as e:
        logger.error(f"Database error getting chat session: {e}")
        return jsonify({"error": "Database connection issue"}), 500
    except Exception as e:
        logger.error(f"API get chat session error: {e}")
        return jsonify({"error": "Failed to fetch chat session"}), 500

@app.route('/api/user/articles')
@login_required
def api_user_articles():
    """Get user's articles via API"""
    try:
        articles = Article.query.filter_by(user_id=current_user.id)\
                               .order_by(Article.created_at.desc()).all()
        return jsonify([article.to_dict() for article in articles])
    except (OperationalError, DatabaseError) as e:
        logger.error(f"Database error in API articles: {e}")
        return jsonify({"error": "Database connection issue"}), 500
    except Exception as e:
        logger.error(f"API articles error: {e}")
        return jsonify({"error": "Failed to fetch articles"}), 500

@app.route('/api/user/stats')
@login_required
def api_user_stats():
    """Get user statistics"""
    try:
        total_articles = Article.query.filter_by(user_id=current_user.id).count()
        total_words = db.session.query(db.func.sum(Article.word_count))\
                               .filter_by(user_id=current_user.id).scalar() or 0
        total_downloads = db.session.query(db.func.sum(Article.download_count))\
                                    .filter_by(user_id=current_user.id).scalar() or 0

        return jsonify({
            'total_articles': total_articles,
            'total_words': total_words,
            'total_downloads': total_downloads,
            'words_this_month': getattr(current_user, 'words_generated_this_month', 0) or 0,
            'downloads_this_month': getattr(current_user, 'downloads_this_month', 0) or 0,
            'word_limit': MONTHLY_WORD_LIMIT,
            'download_limit': MONTHLY_DOWNLOAD_LIMIT,
            'member_since': current_user.created_at.isoformat() if current_user.created_at else None
        })
    except (OperationalError, DatabaseError) as e:
        logger.error(f"Database error in API stats: {e}")
        return jsonify({"error": "Database connection issue"}), 500
    except Exception as e:
        logger.error(f"API stats error: {e}")
        return jsonify({"error": "Failed to fetch stats"}), 500

@app.route('/api/articles/<int:article_id>/download', methods=['POST'])
@login_required
def api_download_article(article_id):
    """API endpoint to download an article"""
    try:
        article = Article.query.filter_by(id=article_id, user_id=current_user.id).first_or_404()

        # Check monthly download quota
        if not check_monthly_download_quota(current_user):
            return jsonify({"error": f"You've reached your monthly limit of {MONTHLY_DOWNLOAD_LIMIT} downloads."}), 403

        # Generate DOCX
        soup = BeautifulSoup(article.content_html, 'html.parser')
        doc = Document()
        doc.add_heading(article.title, level=0)

        for element in soup.find_all(['h2', 'h3', 'p', 'div']):
            if element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'p' and not element.find_parents("div"):
                doc.add_paragraph(element.get_text())
            elif element.name == 'div' and "real-image-container" in element.get('class', []):
                title_p = element.find('p', class_='image-title')
                img_tag = element.find('img')
                alt_text_p = element.find('p', class_='alt-text-display')
                attr_p = element.find('p', class_='attribution')

                if title_p:
                    p = doc.add_paragraph(title_p.get_text())
                    p.alignment = 1
                    p.bold = True

                if img_tag and img_tag.get('src'):
                    try:
                        img_response = requests.get(img_tag['src'], stream=True, timeout=10)
                        img_response.raise_for_status()
                        doc.add_picture(io.BytesIO(img_response.content), width=Inches(5.5))
                    except requests.RequestException:
                        doc.add_paragraph(f"[Image failed to load from {img_tag['src']}]")

                if alt_text_p:
                    p = doc.add_paragraph()
                    clean_alt_text = alt_text_p.get_text()
                    if clean_alt_text.lower().startswith("alt text:"):
                        clean_alt_text = clean_alt_text[len("Alt Text:"):].strip()
                    run = p.add_run(clean_alt_text)
                    run.italic = True
                    p.alignment = 1

                if attr_p:
                    p = doc.add_paragraph(attr_p.get_text())
                    p.alignment = 1
                    p.italic = True

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        filename = f"{article.title[:50].strip().replace(' ', '_')}.docx"

        # Update download count
        article.increment_download()
        current_user.downloads_this_month = (getattr(current_user, 'downloads_this_month', 0) or 0) + 1
        db.session.commit()

        return send_file(file_stream, as_attachment=True, download_name=filename,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except (OperationalError, DatabaseError) as e:
        logger.error(f"Database error in API download: {e}")
        return jsonify({"error": "Database connection issue"}), 500
    except Exception as e:
        logger.error(f"API download error: {e}")
        return jsonify({"error": "Failed to download article"}), 500

@app.route('/api/v1/studio/stats/<studio_name>')
@login_required
def get_studio_stats(studio_name):
    """Get usage statistics for a specific studio."""
    if studio_name not in STUDIO_TYPE_MAPPING:
        return jsonify({"error": "Invalid studio name"}), 404

    studio_types = STUDIO_TYPE_MAPPING[studio_name]

    # Get count for the current month
    start_of_month = datetime.utcnow().replace(day=1, hour=0, minute=0, second=0, microsecond=0)

    count = db.session.query(func.count(ChatSession.id))\
                      .filter(ChatSession.user_id == current_user.id,
                              ChatSession.studio_type.in_(studio_types),
                              ChatSession.created_at >= start_of_month)\
                      .scalar()

    # For now, we'll just return the count. We can expand this later.
    stats = {
        f"{studio_name}_usage_this_month": count,
        "monthly_word_limit": MONTHLY_WORD_LIMIT, # This is a global limit for now
        "words_this_month": getattr(current_user, 'words_generated_this_month', 0) or 0
    }

    return jsonify(stats)

@app.route('/api/chat-sessions/<string:session_id>', methods=['DELETE'])
@login_required
def delete_chat_session(session_id):
    """Delete a chat session and all associated articles"""
    try:
        # Find the chat session
        chat_session = ChatSession.query.filter_by(session_id=session_id, user_id=current_user.id).first()
        if not chat_session:
            return jsonify({"error": "Chat session not found"}), 404

        # Delete all articles associated with this chat session
        Article.query.filter_by(chat_session_id=chat_session.id, user_id=current_user.id).delete()

        # Delete the chat session
        db.session.delete(chat_session)
        db.session.commit()

        return jsonify({"success": True, "message": "Chat session and associated articles deleted"})
    except (OperationalError, DatabaseError) as e:
        db.session.rollback()
        logger.error(f"Database error deleting chat session: {e}")
        return jsonify({"error": "Database connection issue"}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Chat session deletion error: {e}")
        return jsonify({"error": "Failed to delete chat session"}), 500

# --- 6. LEGAL PAGES ---

@app.route('/privacy')
def privacy_policy():
    """Privacy Policy page"""
    return render_template('legal/privacy.html')

@app.route('/terms')
def terms_of_service():
    """Terms of Service page"""
    return render_template('legal/terms.html')

@app.route('/support')
def support():
    """Support page"""
    return render_template('legal/support.html')

@app.route('/contact', methods=['GET', 'POST'])
def contact():
    """Contact page"""
    if request.method == 'POST':
        try:
            name = request.form.get('name', '').strip()
            email = request.form.get('email', '').strip()
            subject = request.form.get('subject', '').strip()
            message = request.form.get('message', '').strip()

            if not all([name, email, subject, message]):
                flash('Please fill in all required fields.', 'error')
                return render_template('legal/contact.html')

            # Send email
            if send_contact_email(name, email, subject, message):
                flash('Thank you for your message! We\'ll get back to you soon.', 'success')
            else:
                flash('Sorry, there was an error sending your message. Please try again later.', 'error')

        except Exception as e:
            logger.error(f"Contact form error: {e}")
            flash('Sorry, there was an error processing your request.', 'error')

    return render_template('legal/contact.html')

# --- 7. ERROR HANDLERS ---

@app.errorhandler(404)
def not_found_error(error):
    return render_template('errors/404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    logger.error(f"Internal server error: {error}")
    return render_template('errors/500.html'), 500

@app.errorhandler(OperationalError)
def database_error(error):
    db.session.rollback()
    logger.error(f"Database operational error: {error}")
    flash('Database connection issue. Please try again in a moment.', 'error')
    return redirect(url_for('index'))

# --- 8. DATABASE INITIALIZATION ---

def init_db():
    """Initialize database tables"""
    try:
        with app.app_context():
            db.create_all()
            logger.info("Database tables created successfully")

            # Run migrations to ensure schema is up to date
            try:
                from migrations import run_migrations
                run_migrations()
            except ImportError:
                logger.warning("Migrations module not found, skipping migrations")
            except Exception as e:
                logger.warning(f"Migration error (non-fatal): {e}")

    except Exception as e:
        logger.error(f"Database initialization error: {e}")

if __name__ == "__main__":
    init_db()
    port = int(os.environ.get('PORT', 5001))
    app.run(host='0.0.0.0', port=port, debug=False)
else:
    # For production deployment
    init_db()
